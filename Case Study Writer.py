import streamlit as st
import google.generativeai as genai
import os
import docx
import uuid
from datetime import datetime
from fpdf import FPDF
import re

GOOGLE_API_KEY = AIzaSyD1WeVGIKaV1oyYlDsk2a_EDWdsSclwMqU
# Configure Gemini API
try:
    genai.configure(api_key=os.environ.get("GOOGLE_API_KEY"))
    model = genai.GenerativeModel('gemini-1.5-flash')
except Exception as e:
    st.error(f"Error configuring Gemini API: {e}")
    st.stop()

# Function to generate case study using Gemini with specified structure
def generate_case_study(introduction, business_need, solution, client, industry, offering, aws_services, results, tone, length):
    try:
        # Define length parameters
        length_prompt = {
            "Short": "in 300-500 words",
            "Medium": "in 500-800 words",
            "Long": "in 800-1200 words"
        }[length]

        # Craft prompt for Gemini with specified structure
        prompt = f"""
        As an experienced Chief Marketing Officer, craft a professional case study with a {tone.lower()} tone, {length_prompt}, optimized for SEO and infused with a human touch. Use the following information and structure it into the specified sections: Introduction, Business Need, Solution Approach and Implementation, Client, Industry, Offering, AWS Services, and Reaping Rewards. Ensure a captivating narrative that resonates with human readers and search engines, weaving data-driven insights with engaging storytelling. Avoid placeholders like [Your Name]. If specific data is missing, make reasonable assumptions to enhance the narrative while maintaining authenticity.

        **Introduction**: {introduction}
        **Business Need**: {business_need}
        **Solution Approach and Implementation**: {solution}
        **Client**: {client}
        **Industry**: {industry}
        **Offering**: {offering}
        **AWS Services**: {aws_services}
        **Reaping Rewards**: {results}

        Format the output in markdown with clear section headers (e.g., ## Introduction). Include at least two specific, realistic data points per section where applicable to enhance credibility.
        """

        # Call Gemini API
        response = model.generate_content(prompt)
        if not response or not response.text:
            raise ValueError("Empty response from Gemini API")
        return response.text
    except Exception as e:
        st.error(f"Error generating case study: {e}")
        return None

# Function to create Word document
def create_word_doc(case_study, title):
    doc = docx.Document()
    doc.add_heading(title, 0)
    sections = re.split(r'##\s+', case_study)[1:]  # Split by markdown headers
    for section in sections:
        section_title, content = section.split('\n', 1)
        doc.add_heading(section_title, level=1)
        doc.add_paragraph(content.strip())
    file_name = f"case_study_{uuid.uuid4()}.docx"
    doc.save(file_name)
    return file_name

# Function to create PDF document using LaTeX
def create_pdf_latex(case_study, title):
    latex_content = r"""
    \documentclass{article}
    \usepackage[utf8]{inputenc}
    \usepackage{geometry}
    \geometry{a4paper, margin=1in}
    \usepackage{parskip}
    \usepackage{titlesec}
    \titleformat{\section}{\Large\bfseries}{\thesection}{1em}{}
    \titleformat{\subsection}{\large\bfseries}{\thesubsection}{1em}{}
    \usepackage{noto}

    \begin{document}
    \begin{center}
        \textbf{\Large """ + title + r"""}
    \end{center}
    \vspace{1cm}
    """ + '\n\n'.join(
        [r"\section{" + section.split('\n', 1)[0].replace('## ', '') + r"}" + section.split('\n', 1)[1].replace('#', r'\#')
         for section in re.split(r'##\s+', case_study)[1:]]
    ) + r"""
    \end{document}
    """

    file_name = f"case_study_{uuid.uuid4()}.tex"
    with open(file_name, 'w', encoding='utf-8') as f:
        f.write(latex_content)
    return file_name

# Streamlit app
st.set_page_config(page_title="Case Study Writer Tool", layout="wide")
st.title("Case Study Writer Tool")
st.write("Enter details to generate a professional, SEO-optimized case study with a human touch, following the specified structure: Introduction, Business Need, Solution Approach and Implementation, Client, Industry, Offering, AWS Services, Reaping Rewards.")

# Input form
with st.form("case_study_form"):
    st.subheader("Input Details")
    col1, col2 = st.columns(2)
    
    with col1:
        introduction = st.text_area("Introduction", placeholder="Introduce the core topic or problem to hook the reader.", height=100)
        business_need = st.text_area("Business Need", placeholder="Outline the specific business challenge or opportunity.", height=100)
        solution = st.text_area("Solution Approach and Implementation", placeholder="Detail the strategy and execution steps.", height=100)
        client = st.text_area("Client", placeholder="Describe the client or company (e.g., size, background).", height=100)
    
    with col2:
        industry = st.text_area("Industry", placeholder="Specify the industry and its challenges/trends.", height=100)
        offering = st.text_area("Offering", placeholder="Explain the product or service and its value proposition.", height=100)
        aws_services = st.text_area("AWS Services (Optional)", placeholder="List any AWS services used, if applicable.", height=100)
        results = st.text_area("Reaping Rewards", placeholder="Summarize outcomes and benefits achieved.", height=100)
    
    tone = st.selectbox("Tone", ["Professional", "Engaging", "Technical", "Narrative"])
    length = st.selectbox("Length", ["Short", "Medium", "Long"])
    
    submitted = st.form_submit_button("Generate Case Study")

# Real-time preview
if submitted:
    if not all([introduction, business_need, solution, client, industry, offering, results]):
        st.error("Please fill in all required fields (AWS Services is optional).")
    else:
        with st.spinner("Generating case study..."):
            case_study = generate_case_study(
                introduction, business_need, solution, client, industry, offering, aws_services, results, tone, length
            )
            if case_study:
                st.subheader("Case Study Preview")
                st.markdown(case_study)
                
                # Store case study in session state for export
                st.session_state['case_study'] = case_study
                st.session_state['title'] = "Case Study Generated on " + datetime.now().strftime("%Y-%m-%d")

# Export options
if 'case_study' in st.session_state:
    st.subheader("Export Options")
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("Download as PDF"):
            with st.spinner("Generating PDF..."):
                latex_file = create_pdf_latex(st.session_state['case_study'], st.session_state['title'])
                # Note: LaTeX compilation requires latexmk, handled by the system
                pdf_file = latex_file.replace('.tex', '.pdf')
                try:
                    with open(pdf_file, "rb") as f:
                        st.download_button(
                            label="Download PDF",
                            data=f,
                            file_name=pdf_file,
                            mime="application/pdf"
                        )
                except FileNotFoundError:
                    st.error("PDF generation requires LaTeX environment (latexmk). Please ensure it's installed or use Word export.")
    
    with col2:
        if st.button("Download as Word"):
            with st.spinner("Generating Word document..."):
                docx_file = create_word_doc(st.session_state['case_study'], st.session_state['title'])
                with open(docx_file, "rb") as f:
                    st.download_button(
                        label="Download Word",
                        data=f,
                        file_name=docx_file,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

# Suggestions for enhancing output quality
st.subheader("Tips for Best Results")
st.markdown("""
- **Detailed Inputs**: Provide specific details for each section to ensure a rich, authentic narrative. For example, include metrics like 'reduced costs by 30%' in Reaping Rewards.
- **SEO Optimization**: Use industry-specific keywords in the Industry and Offering sections to enhance searchability.
- **Human Touch**: Add relatable context in the Client section, such as a brief story about their goals or challenges.
- **AWS Integration**: If using AWS Services, specify tools like 'AWS Lambda' or 'Amazon S3' to add technical credibility.
- **Iterate and Refine**: If the output needs adjustment, tweak inputs or tone to align with your audience (e.g., Narrative for storytelling, Technical for industry reports).
- **Validate Outputs**: Review the generated case study for accuracy, especially for assumed data points, to ensure alignment with your brand’s voice.
""")
